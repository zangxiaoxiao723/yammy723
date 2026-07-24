from __future__ import annotations

import csv
from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

import build_hpdi_customer_word_v2 as word


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs" / "hpdi_followup_v3"
PLOTS = OUT / "plots"
DOCX = OUT / "HPDI低温泵0716-0723后续条件声音对比报告.docx"


def add_table(doc, headers, rows, widths, center_cols=()):
    table = doc.add_table(rows=1, cols=len(headers))
    for i, text in enumerate(headers):
        word.shade(table.cell(0, i), word.DARK_BLUE)
        word.set_cell_text(
            table.cell(0, i), text, bold=True, color=word.WHITE,
            align=WD_ALIGN_PARAGRAPH.CENTER, size=8.7,
        )
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            word.set_cell_text(
                cells[i], text, bold=(i == 0), size=8.5,
                align=WD_ALIGN_PARAGRAPH.CENTER if i in center_cols else WD_ALIGN_PARAGRAPH.LEFT,
            )
    word.set_table_geometry(table, widths)
    return table


def add_compact_figure(doc, path, caption, width=5.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(6)
    r = cap.add_run(caption)
    word.set_run_font(r, size=9, color="555555", italic=True)


def short_condition(series_id: str) -> str:
    labels = {
        "FR_0716_CYL_800": "嵌套油缸",
        "FR_0716_CYL_900": "嵌套油缸",
        "FR_0716_CYL_1000": "嵌套油缸",
        "FR_0716_CYL_1125": "嵌套油缸",
        "FR_0716_SHUTTLE_900_R1": "嵌套油缸+西港梭阀 R1",
        "FR_0716_SHUTTLE_900_R2": "嵌套油缸+西港梭阀 R2",
        "FR_0716_SHUTTLE_1000": "嵌套油缸+西港梭阀",
        "FR_0722_B_HEAT_900": "热处理活塞+B样",
        "FR_0722_B_HEAT_1000": "热处理活塞+B样",
        "FR_DELIVERY_0723_700": "0723交付泵",
    }
    return labels[series_id]


def judgment(row):
    sid = row["series_id"]
    if sid == "FR_DELIVERY_0723_700":
        return "重咚明显收敛"
    if sid == "FR_0716_SHUTTLE_900_R2":
        return "与R1不一致，需复测"
    if sid == "FR_0722_B_HEAT_900":
        return "较响声未改善"
    if row["dd_loud_delta_db"]:
        delta = float(row["dd_loud_delta_db"])
        return "较响声改善" if delta <= -2 else "接近基准" if delta <= 0 else "较响声偏高"
    delta = float(row["initial_loud_delta_db"])
    return "较初始改善" if delta <= -2 else "较初始接近" if delta <= 1 else "较初始偏高"


def main():
    with (OUT / "followup_comparison.csv").open(encoding="utf-8-sig") as f:
        rows = list(csv.DictReader(f))
    by_id = {r["series_id"]: r for r in rows}

    doc = word.setup_document()
    header = doc.sections[0].header.paragraphs[0]
    header.clear()
    word.set_run_font(header.add_run("HPDI低温泵声音工程对比 | 0716-0723后续条件"), size=8.5, color="64748B")

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(8)
    title.paragraph_format.space_after = Pt(2)
    word.set_run_font(title.add_run("后续条件声音对比报告"), size=24, bold=True, color=word.INK)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    word.set_run_font(subtitle.add_run("HPDI低温泵 | 0716、0722、0723测试"), size=13, bold=True, color=word.TEAL)

    meta = doc.add_table(rows=4, cols=2)
    word.set_table_geometry(meta, [1800, 7560])
    metadata = [
        ("测试范围", "10段后续视频；700/800/900/1000/1125 rpm"),
        ("对比基准", "东德同转速基准优先；无东德视频时仅与富瑞初始比较"),
        ("主要指标", "整体咚声趋势、常见咚声、较响咚声"),
        ("报告版本", "V3 | 2026-07-24"),
    ]
    for i, (label, value) in enumerate(metadata):
        word.shade(meta.cell(i, 0), word.LIGHT_GRAY)
        word.set_cell_text(meta.cell(i, 0), label, bold=True, color=word.DARK_BLUE)
        word.set_cell_text(meta.cell(i, 1), value)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    word.add_callout(
        doc, "核心结论",
        "0723交付泵的主观改善有数据支持：700 rpm较响咚声比东德低2.1 dB、比富瑞初始低4.7 dB，且本段没有咚声超过东德的较响门槛。它的常见咚声仍比东德高3.4 dB，因此结论应表述为“重咚收敛、整体听感改善”，而不是“所有咚声都更小”。",
        word.GOLD,
    )

    doc.add_page_break()
    word.add_heading(doc, "1. 本批测试结果总览", 1)
    overview = []
    for row in rows:
        overview.append([
            row["speed_rpm"], short_condition(row["series_id"]), row["event_count"],
            row["vs_dd_common"], row["vs_dd_loud"], row["vs_initial_loud"], judgment(row),
        ])
    add_table(
        doc,
        ["转速", "测试条件", "有效咚声", "常见声/东德", "较响声/东德", "较响声/初始", "判断"],
        overview, [720, 2040, 900, 1320, 1320, 1320, 1740], center_cols=(0, 2, 3, 4, 5, 6),
    )
    word.add_para(doc, "表中“低于”表示该组声音更小，“高于”表示该组声音更大。1000 rpm没有东德同转速视频，因此对应单元格显示“无同转速基准”，只看相对富瑞初始的变化。", after=4)
    word.add_callout(doc, "整体判断", "0716嵌套油缸在800、900、1000 rpm的较响咚声均较初始下降；0722热处理活塞+B样在900 rpm反而偏高；西港梭阀两次900 rpm结果差异大，当前不能据此确认稳定改善。", word.LIGHT_BLUE)

    doc.add_page_break()
    word.add_heading(doc, "2. 0723交付泵：为什么听起来更好", 1)
    delivery = by_id["FR_DELIVERY_0723_700"]
    add_table(doc, ["指标", "相对东德700 rpm", "相对富瑞初始700 rpm", "含义"], [[
        "常见咚声", delivery["vs_dd_common"], delivery["vs_initial_common"], "多数普通周期仍偏响",
    ], [
        "较响咚声", delivery["vs_dd_loud"], delivery["vs_initial_loud"], "最影响听感的重咚明显降低",
    ], [
        "超过东德较响门槛", "0.0%", "不适用", "本段未出现高于东德门槛的重咚",
    ]], [1800, 2100, 2200, 3260], center_cols=(0, 1, 2))
    word.add_figure(doc, PLOTS / "04_交付泵700rpm详细对比.png", "图1  700 rpm常见咚声与较响咚声的单次曲线对比")
    word.add_callout(doc, "客户沟通建议", "可以说明“交付泵已明显压低偶发重咚，较响咚声优于东德约2.1 dB；普通周期的常见声仍有继续优化空间”。这比笼统说“声音更小”更符合数据。", word.GREEN)

    doc.add_page_break()
    word.add_heading(doc, "3. 0716嵌套油缸：较响咚声有改善", 1)
    cyl_ids = ["FR_0716_CYL_800", "FR_0716_CYL_900", "FR_0716_CYL_1000", "FR_0716_CYL_1125"]
    cyl_rows = []
    for sid in cyl_ids:
        r = by_id[sid]
        cyl_rows.append([r["speed_rpm"], r["vs_dd_common"], r["vs_dd_loud"], r["vs_initial_loud"], r["above_dd_loud_share_pct"] or "无基准"])
    add_table(doc, ["转速", "常见声/东德", "较响声/东德", "较响声/初始", "超过东德门槛"], cyl_rows, [1000, 1900, 1900, 1900, 2660], center_cols=(0, 1, 2, 3, 4))
    word.add_para(doc, "800 rpm的较响咚声比东德低3.1 dB、比初始低4.1 dB；900 rpm比东德低2.1 dB、比初始低1.2 dB；1000 rpm比初始低2.8 dB。说明更换嵌套油缸后，重咚总体收敛。", bold_label="主要改善：")
    word.add_para(doc, "常见咚声并未同步下降，800、900 rpm仍高于东德。1125 rpm东德录音中的常见声提取值异常偏低，所以该转速只重点参考较响声，不用常见声做客户结论。", bold_label="仍需注意：")

    word.add_heading(doc, "4. 900 rpm不同配置：梭阀重复性不足，0722状态偏响", 1)
    word.add_figure(doc, PLOTS / "05_900rpm各配置对比.png", "图2  900 rpm各配置的常见咚声和较响咚声")
    word.add_para(doc, "西港梭阀R1的较响咚声比东德低2.9 dB，但R2反而比东德高2.3 dB。R2只识别到14次有效咚声，样本短，且与R1方向相反，因此当前应标记为“重复性不足”，不能下稳定改善结论。")
    word.add_para(doc, "0722热处理活塞+B样在900 rpm的常见声比东德高8.1 dB、较响声高2.5 dB；这一状态没有改善。1000 rpm较响声比富瑞初始高0.6 dB，基本接近但没有优势。")

    doc.add_page_break()
    word.add_heading(doc, "5. 整段测试中的声音稳定性", 1)
    word.add_para(doc, "下图每个点代表一次有效咚声，深色线表示连续5次咚声的移动趋势。它用来观察一段测试中是否存在“连续几个周期较小，随后几个周期变大”的现象。")
    add_compact_figure(doc, PLOTS / "01_后续测试整体咚声趋势.png", "图3  10段后续测试的整体咚声趋势", width=4.9)
    word.add_callout(doc, "直观观察", "0723交付泵的主要咚声集中在约73-79 dB参考区间，趋势较平稳；西港梭阀R2和0722两组视频较短，点数较少，后续建议每个稳定工况连续录制至少60秒。", word.LIGHT_BLUE)

    doc.add_page_break()
    word.add_heading(doc, "6. 数据口径：为什么本版改用背景归一化声级", 1)
    word.add_para(doc, "新批视频与7月初基准视频跨日期录制。虽然手机、角度和距离保持一致，但手机自动增益会随现场背景变化，新视频的背景录音电平整体比基准高约7-11 dB。若直接沿用固定分贝偏移，会把所有新测试误判为整体变响。")
    word.add_para(doc, "本版先在每段视频中识别泵咚声附近的局部台架背景，再把该段背景统一平移到65 dB参考线，然后比较泵咚声高出自身背景的程度。图表纵轴因此标为“背景归一化声级(dB)”。数值是正数，差值写成“高于/低于X.X dB”。")
    method_rows = [
        ["泵咚声提取", "45-650 Hz，保留往复冲击的低频主体"],
        ["台架嗡声处理", "每次咚声使用前后局部背景做功率扣除"],
        ["排气尖声处理", "监测1800-9000 Hz，高频占比异常的时段不计入咚声统计"],
        ["常见咚声", "表示一段视频中大多数普通工作周期的典型水平"],
        ["较响咚声", "表示一段视频中声音较大的约一成工作周期"],
        ["正式比较", "优先同转速；不同转速只观察趋势，不直接判定优劣"],
    ]
    add_table(doc, ["处理环节", "本版做法"], method_rows, [2100, 7260])
    word.add_callout(doc, "使用边界", "背景归一化可以显著降低手机自动增益和台架稳态背景的影响，但不能替代1级声级计的同步定点测量。当前结果适合产品调试和同条件工程对比，不作为法规或认证声级。", word.GOLD)

    word.add_heading(doc, "7. 下一轮测试建议", 1)
    recommendations = [
        ["1", "同一转速至少重复2次，每段稳定工作不少于60秒"],
        ["2", "优先复测西港梭阀900 rpm，确认R1/R2差异是否可重复"],
        ["3", "交付泵补测800、900、1000 rpm，才能形成完整同转速基准对比"],
        ["4", "手机固定手动增益时关闭自动增益；若无法关闭，继续使用本版背景归一化口径"],
        ["5", "条件允许时同步记录换向信号、油压和低温端压力，用于定位重咚对应动作"],
    ]
    add_table(doc, ["序号", "要求"], recommendations, [900, 8460], center_cols=(0,))

    doc.save(DOCX)
    print(DOCX)


if __name__ == "__main__":
    main()
