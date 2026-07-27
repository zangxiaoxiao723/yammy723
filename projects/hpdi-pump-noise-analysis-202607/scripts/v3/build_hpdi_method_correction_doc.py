from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

import build_hpdi_customer_word_v2 as word


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs" / "hpdi_followup_v3"
DOCX = OUT / "HPDI后续声音分析口径复核与修正说明_内部.docx"


def table(doc, headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    for i, text in enumerate(headers):
        word.shade(t.cell(0, i), word.DARK_BLUE)
        word.set_cell_text(t.cell(0, i), text, bold=True, color=word.WHITE, align=WD_ALIGN_PARAGRAPH.CENTER, size=9)
    for row in rows:
        cells = t.add_row().cells
        for i, text in enumerate(row):
            word.set_cell_text(cells[i], text, align=WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT, size=9)
    word.set_table_geometry(t, widths)
    return t


def main():
    doc = word.setup_document()
    header = doc.sections[0].header.paragraphs[0]
    header.clear()
    word.set_run_font(header.add_run("HPDI低温泵声音分析 | 口径复核与修正"), size=8.5, color="64748B")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    word.set_run_font(p.add_run("后续声音分析口径复核与修正说明"), size=22, bold=True, color=word.INK)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    word.set_run_font(p.add_run("涉及0716-0723后续测试及可选对比工具"), size=12.5, bold=True, color=word.TEAL)

    word.add_callout(doc, "修正结论", "V3曾把每段背景对齐后的“咚声突出度”作为主声级，并反向重算东德和富瑞初始基准。这违反了基准锁定原则。现已撤销该做法：V2基准数值、算法和固定校准偏移全部恢复，后续测试只能按同一固定方法追加。", word.RED)

    word.add_heading(doc, "1. 用户发现的问题", 1)
    word.add_para(doc, "同样选择东德900 rpm和富瑞初始900 rpm，上一版工具显示富瑞初始常见咚声高于东德7.6 dB、较响咚声高于东德5.1 dB；V3错误版本却显示高1.6 dB和低0.9 dB。该差异不是Excel引用错误，而是分析口径被改变。")
    table(doc, ["900 rpm项目", "V2既定基准", "V3错误归一化", "变化原因"], [
        ["东德背景平移", "0.0 dB", "+3.5 dB", "背景被强制对齐到65 dB"],
        ["富瑞初始背景平移", "0.0 dB", "-2.5 dB", "背景被强制对齐到65 dB"],
        ["常见咚声差", "富瑞高7.6 dB", "富瑞高1.6 dB", "相对差被压缩约6.0 dB"],
        ["较响咚声差", "富瑞高5.1 dB", "富瑞低0.9 dB", "结论方向被错误翻转"],
    ], [2400, 1800, 1900, 3260])

    word.add_heading(doc, "2. 为什么背景对齐不能覆盖基准", 1)
    word.add_para(doc, "V2固定校准偏移来自富瑞初始1000 rpm视频中的分贝仪画面，并统一用于基准视频。它不是每个视频独立校准，但一旦被确定为项目基准，后续对比必须保持算法不变。")
    word.add_para(doc, "每段背景对齐实际比较的是“咚声相对本段背景有多突出”。背景变化既可能来自手机自动增益，也可能来自真实台架嗡声、泵连续噪声或工况变化。没有独立证据时，不能把全部背景差都当成手机增益，更不能用该平移值修改既定基准。")
    doc.add_page_break()
    word.add_callout(doc, "正确命名", "背景对齐结果只能作为“咚声突出度辅助指标”，不能称为绝对声级、背景归一化声级主结论，也不能替代固定基准结果。", word.GOLD)

    word.add_heading(doc, "3. 对0723交付泵结论的影响", 1)
    table(doc, ["比较方法", "较响咚声相对东德700", "能否作为声级结论"], [
        ["V2固定基准算法", "高于8.0 dB", "受跨日期手机自动增益影响，需复测"],
        ["背景突出度辅助指标", "低于2.1 dB", "只能说明重咚相对本段背景不突出"],
    ], [3000, 2600, 3760])
    word.add_para(doc, "两种方法方向相反，说明现有0723无分贝仪视频不足以量化证明交付泵绝对声音优于东德。用户主观感觉“交付泵更好”仍可作为试验观察，但不能用当前软件结果对客户宣称较响声低2.1 dB。", bold_label="正式修正：")

    word.add_heading(doc, "4. 修正后的基准管理规则", 1)
    rules = [
        ["1", "东德和富瑞初始V2数值永久锁定；任何后续脚本不得覆盖。"],
        ["2", "主对比固定使用V2频带、事件提取、背景功率扣除和校准偏移。"],
        ["3", "新增算法只能作为独立辅助列或独立工作表，不能改写历史基准。"],
        ["4", "无分贝仪的跨日期视频必须标记手机自动增益风险，不直接形成客户声级结论。"],
        ["5", "后续正式测试应同步声级计，或锁定手机手动录音增益，并保留开始/结束校准声。"],
    ]
    table(doc, ["序号", "固定要求"], rules, [900, 8460])

    word.add_heading(doc, "5. 已完成的修正", 1)
    word.add_para(doc, "已生成《HPDI泵声音固定基准对比工具_复核修正版.xlsx》。工具默认打开东德900和富瑞初始900，结果恢复为常见声高7.6 dB、较响声高5.1 dB；另设“既定基准结果”工作表用于核对全部V2基准。")
    word.add_para(doc, "0716-0723后续数据仍保留在工具中，但统一标记为固定算法下的工程估算。背景突出度结果不再参与主声级判断。")

    doc.save(DOCX)
    print(DOCX)


if __name__ == "__main__":
    main()
