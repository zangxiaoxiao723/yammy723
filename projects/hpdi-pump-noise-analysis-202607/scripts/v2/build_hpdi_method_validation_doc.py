from __future__ import annotations

import csv
from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

from build_hpdi_customer_word_v2 import (
    DARK_BLUE,
    GOLD,
    GREEN,
    INK,
    LIGHT_BLUE,
    LIGHT_GRAY,
    RED,
    TEAL,
    WHITE,
    add_callout,
    add_figure,
    add_heading,
    add_para,
    set_cell_text,
    set_run_font,
    set_table_geometry,
    setup_document,
    shade,
)


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs" / "hpdi_baseline_v2"
PLOTS = OUT / "plots"
DOCX = OUT / "HPDI声音分析方法准确性与可复现性说明.docx"


def main() -> None:
    with (OUT / "method_consistency_summary.csv").open(encoding="utf-8-sig") as f:
        consistency = list(csv.DictReader(f))

    doc = setup_document()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Inches(0.1)
    p.paragraph_format.space_after = Inches(0.03)
    r = p.add_run("声音分析方法准确性与可复现性说明")
    set_run_font(r, size=23, bold=True, color=INK)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Inches(0.18)
    r = p2.add_run("HPDI低温泵 | 干扰排除、包络提取与基准比较证据链")
    set_run_font(r, size=12.5, bold=True, color=TEAL)

    add_callout(doc, "先说结论", "软件名称本身不能证明结果正确。可信度来自固定输入、固定算法、多方法交叉验证、统计区间、原始文件校验和重复试验。当前结果可用于同条件工程比较，但尚不能替代认证声级测试。", GOLD)

    add_heading(doc, "1. 当前结果能证明到什么程度", 1)
    add_para(doc, "算法可复现性已经具备：同一 WAV 输入、同一脚本和参数会得到同一事件、同一曲线和同一比较结果；源文件通过 SHA-256 校验，能证明分析过程中没有更换录音。")
    add_para(doc, "单次测试内部统计已经具备：每段录音包含几十次咚声，可计算常见水平、较响水平和事件重采样区间。")
    add_para(doc, "试验重复性尚未完全具备：目前每个产品、每个转速主要只有一段有效录音。几十次咚声属于同一次试验，不能替代拆装后重复、不同日期重复和测试顺序随机化。")
    add_callout(doc, "对领导的准确表述", "现有分析能够证明某些转速下的声音差异具有较强内部一致性，但不能宣称所有绝对 dB(A)均为认证值，也不能用单次录音证明长期量产一致性。", LIGHT_BLUE)

    add_heading(doc, "2. 使用的软件和可追溯文件", 1)
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [1900, 2300, 5160])
    for i, text in enumerate(["工具", "版本", "用途"]):
        shade(table.cell(0, i), DARK_BLUE)
        set_cell_text(table.cell(0, i), text, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    tools = [
        ("FFmpeg", "8.0.1", "从 MP4 无损提取单声道 44.1 kHz PCM 音频。"),
        ("Python", "3.11.9", "执行固定分析脚本和批量处理。"),
        ("SciPy", "1.17.1", "数字滤波、A计权、短时能量、峰值检测。"),
        ("NumPy", "2.0.2", "功率计算、统计分位和事件重采样。"),
        ("Matplotlib", "3.10.8", "输出包络、单声事件和一致性验证图。"),
        ("AR824分贝仪", "视频可见 dBA", "提供现场总声级参考；FAST/SLOW标志未能从画面确认。"),
    ]
    for item in tools:
        cells = table.add_row().cells
        for i, text in enumerate(item):
            set_cell_text(cells[i], text, bold=(i == 0), align=WD_ALIGN_PARAGRAPH.CENTER if i < 2 else WD_ALIGN_PARAGRAPH.LEFT, size=9)
    add_para(doc, "交付数据中保存源 WAV 路径和 SHA-256；保存每次咚声的时间、声级、背景和污染指标；保存汇总 CSV、分析脚本和图表。任何人可从事件明细追溯到原录音时间点。")

    doc.add_page_break()
    add_heading(doc, "3. 从原视频到比较结果的完整流程", 1)
    flow = doc.add_table(rows=6, cols=3)
    set_table_geometry(flow, [900, 2200, 6260])
    for i, text in enumerate(["步骤", "处理", "固定口径"]):
        shade(flow.cell(0, i), DARK_BLUE)
        set_cell_text(flow.cell(0, i), text, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    steps = [
        ("1", "提取音频", "MP4 转为单声道 44.1 kHz PCM，去除直流分量，不做人工音量归一化。"),
        ("2", "限制频带", "泵咚声主通道 45-650 Hz；排气污染监测通道 1800-9000 Hz。"),
        ("3", "A计权", "按人耳频率敏感度计算估算 dB(A)；同时保留未计权结果做交叉检查。"),
        ("4", "提取包络", "80 ms 短时均方根窗口、20 ms 步长；显示连续冲击强弱变化。"),
        ("5", "检测咚声", "按包络突出度检测，相邻事件至少 0.18 s；高频排气占优的峰值剔除。"),
        ("6", "计算比较", "事件窗口为峰值前 60 ms 至后 180 ms；扣除前后局部背景功率；只比较相同转速。"),
    ]
    for step in steps:
        cells = flow.add_row().cells
        for i, text in enumerate(step):
            set_cell_text(cells[i], text, bold=(i == 0), align=WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT, size=9)

    add_heading(doc, "4. 如何排除台架嗡声和排气尖声", 1)
    add_para(doc, "台架嗡声：单独录制的“台架1/台架2”用于识别稳定频谱。正式事件不直接使用整段总声压，而是在咚声前后各取局部背景，在功率域相减。稳定持续的嗡声因此被大幅抵消。", bold_label="台架嗡声：")
    add_para(doc, "排气尖声：排气主要集中在高频。算法单独计算 1800-9000 Hz 包络；若高频包络相对 45-650 Hz 泵冲击通道异常占优，该事件标记为污染并不进入咚声统计。", bold_label="排气尖声：")
    add_para(doc, "防止误删：排气高频通道只用于判断污染，不直接决定泵咚声大小；泵声指标始终来自低中频冲击通道。", bold_label="防止误删：")
    add_callout(doc, "边界", "如果排气声同时包含很强的低频脉冲，并且与泵咚声在同一时刻重叠，仅靠单通道手机录音无法百分之百分离。这类区间只能剔除或在复测时避开。", RED)

    add_heading(doc, "5. 包络和“常见/较响”如何定义", 1)
    defs = doc.add_table(rows=3, cols=2)
    set_table_geometry(defs, [2200, 7160])
    for i, (label, text) in enumerate([
        ("整体包络", "每 20 ms 更新一次 80 ms 短时能量，用于看连续若干周期由小变大或由大变小。"),
        ("常见咚声", "把有效咚声按大小排列，取中间位置，代表大多数正常周期的典型听感。"),
        ("较响咚声", "取偏响的约 10%事件，代表客户更容易注意到的冲击，不等同于单个最大值。"),
    ]):
        shade(defs.cell(i, 0), LIGHT_BLUE)
        set_cell_text(defs.cell(i, 0), label, bold=True, color=DARK_BLUE)
        set_cell_text(defs.cell(i, 1), text)
    add_para(doc, "正文不使用 P50/P90/P95 术语，是为了便于客户阅读；算法内部保留这些统计位置，确保每次新增测试使用同一规则。")

    doc.add_page_break()
    add_heading(doc, "6. 多方法交叉验证：避免单一算法决定结论", 1)
    add_para(doc, "同一批咚声同时使用 5 种方法复核：45-650 Hz 未计权、45-650 Hz A计权、45-120 Hz、120-250 Hz 和 250-650 Hz。每种方法对事件随机重采样 4000 次，形成 95%区间。浅色柱表示区间跨过“无差异”，不能下确定结论。")
    add_figure(doc, PLOTS / "06_全部转速多方法一致性.png", "图1  全部匹配转速的多方法一致性验证")

    summary_table = doc.add_table(rows=1, cols=4)
    set_table_geometry(summary_table, [1100, 1600, 3800, 2860])
    for i, text in enumerate(["转速", "指标", "5种方法结果", "证据判断"]):
        shade(summary_table.cell(0, i), DARK_BLUE)
        set_cell_text(summary_table.cell(0, i), text, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER, size=9)
    for row in consistency:
        low = int(row["富瑞较低"])
        high = int(row["富瑞较高"])
        uncertain = int(row["无法区分"])
        if high >= 4:
            judgement, fill = "较强支持富瑞偏响", RED
        elif low >= 4:
            judgement, fill = "较强支持富瑞较低", GREEN
        else:
            judgement, fill = "当前不能定论", GOLD
        cells = summary_table.add_row().cells
        values = [f"{row['speed_rpm']} rpm", row["metric"], f"较低 {low} / 较高 {high} / 无法区分 {uncertain}", judgement]
        for i, text in enumerate(values):
            if i == 3:
                shade(cells[i], fill)
            set_cell_text(cells[i], text, bold=(i in [0, 3]), align=WD_ALIGN_PARAGRAPH.CENTER, size=8.8)

    add_callout(doc, "当前最稳的结论", "900 rpm 和 1125 rpm 的较响咚声得到多方法支持，富瑞偏响；800 rpm 的常见咚声多方法支持较低，但较响咚声不能定论；700 rpm 两类指标均需要重复试验。", LIGHT_BLUE)

    add_heading(doc, "7. 分贝仪校准为什么仍是“估算 dB(A)”", 1)
    add_para(doc, "视频画面可读取 AR824 的 dBA 数值，1000 rpm 样本中可见 93.9、94.5、93.3、95.3、93.8 和 90.3 dBA。手机全频 A计权电平与前 5 个读数可建立约 117.2 dB 的偏移，但最后一个点没有同步下降，说明手机自动增益或仪器时间响应造成动态不一致。")
    add_para(doc, "因此报告用正数的估算 dB(A)帮助阅读包络；真正用于产品判断的是同条件下“高于东德/低于东德多少 dB”。固定偏移在相减时会抵消，故相对差值比绝对声级可靠。")
    add_callout(doc, "不能这样对外表述", "不能说手机录音已经等同于校准声级计，也不能把估算 dB(A)作为法规、认证或合同验收值。", RED)

    doc.add_page_break()
    add_heading(doc, "8. 可信度分级和升级条件", 1)
    grades = doc.add_table(rows=4, cols=3)
    set_table_geometry(grades, [1100, 3500, 4760])
    for i, text in enumerate(["等级", "所需证据", "允许表述"]):
        shade(grades.cell(0, i), DARK_BLUE)
        set_cell_text(grades.cell(0, i), text, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    grade_rows = [
        ("A级", "每个状态至少 3 次独立重复；背景差小于约 1 dB；4/5以上方法同向；95%区间不跨无差异；重复标准差不大于约 1.5 dB。", "可以向客户说明稳定优于/劣于基准。"),
        ("B级", "单次录音内部事件充足；多方法同向；事件重采样区间不跨无差异。", "可以作为工程判断和优化方向，不宣称长期重复性。"),
        ("C级", "方法方向不一致，或统计区间跨无差异，或背景/排气污染明显。", "只能说接近或待复测，不能判定优劣。"),
    ]
    for row, fill in zip(grade_rows, [GREEN, LIGHT_BLUE, GOLD]):
        cells = grades.add_row().cells
        shade(cells[0], fill)
        for i, text in enumerate(row):
            set_cell_text(cells[i], text, bold=(i == 0), align=WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT, size=9)

    add_heading(doc, "9. 建议的下一轮验证试验", 1)
    protocol = [
        "每个产品、每个转速至少录制 3 次，每次稳定工作不少于 60 s；测试顺序交叉或随机，避免温度随时间变化形成偏差。",
        "手机和分贝仪使用固定支架并做位置标记；保持同一录音软件和模式。若能关闭手机自动增益，优先关闭。",
        "开始和结束前记录台架背景；确保同转速两段背景差不超过约 1 dB，否则该组不作为正式比较。",
        "条件允许时增加一支可导出 PCM 的测量麦克风，并用 94 dB、1 kHz 声校准器在测试前后校准。",
        "同步记录换向信号、油缸位移、系统油压和低温泵出口压力；用统一触发信号定位咚声动作阶段。",
    ]
    for i, text in enumerate(protocol, 1):
        add_para(doc, f"{i}. {text}")

    add_heading(doc, "10. 可复核交付物", 1)
    files = [
        ("源文件校验", "source_file_hashes.csv"),
        ("逐事件明细", "event_details.csv"),
        ("全部转速一致性", "all_speed_method_validation.csv"),
        ("一致性汇总", "method_consistency_summary.csv"),
        ("主分析脚本", "build_hpdi_baseline_v2.py"),
        ("验证脚本", "validate_hpdi_all_speeds.py / validate_hpdi_800_methods.py"),
    ]
    file_table = doc.add_table(rows=1, cols=2)
    set_table_geometry(file_table, [2200, 7160])
    for i, text in enumerate(["内容", "文件"]):
        shade(file_table.cell(0, i), DARK_BLUE)
        set_cell_text(file_table.cell(0, i), text, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    for label, name in files:
        cells = file_table.add_row().cells
        set_cell_text(cells[0], label, bold=True)
        set_cell_text(cells[1], name)

    doc.save(DOCX)
    print(DOCX)


if __name__ == "__main__":
    main()
