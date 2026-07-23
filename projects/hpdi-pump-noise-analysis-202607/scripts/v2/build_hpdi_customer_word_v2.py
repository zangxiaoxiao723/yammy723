from __future__ import annotations

import csv
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs" / "hpdi_baseline_v2"
PLOTS = OUT / "plots"
DOCX = OUT / "HPDI低温泵声音对比基准报告_客户版.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "172033"
TEAL = "007C91"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
GREEN = "DCFCE7"
GOLD = "FEF3C7"
RED = "FEE2E2"
WHITE = "FFFFFF"
LINE = "D8DEE8"


def set_run_font(run, size=None, bold=None, color=None, italic=None):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if italic is not None:
        run.italic = italic


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_cell_text(cell, text, bold=False, color=INK, align=WD_ALIGN_PARAGRAPH.LEFT, size=9.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.1
    run = p.add_run(str(text))
    set_run_font(run, size=size, bold=bold, color=color)


def add_para(doc, text, bold_label=None, after=6, color=INK, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if bold_label and text.startswith(bold_label):
        a = p.add_run(bold_label)
        set_run_font(a, size=size, bold=True, color=color)
        b = p.add_run(text[len(bold_label):])
        set_run_font(b, size=size, color=color)
    else:
        run = p.add_run(text)
        set_run_font(run, size=size, color=color)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_callout(doc, label, text, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    shade(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.2
    r1 = p.add_run(label + "  ")
    set_run_font(r1, size=11, bold=True, color=DARK_BLUE)
    r2 = p.add_run(text)
    set_run_font(r2, size=11, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_figure(doc, path, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    run.add_picture(str(path), width=Inches(6.45))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    cap.paragraph_format.keep_with_next = True
    r = cap.add_run(caption)
    set_run_font(r, size=9, color="555555", italic=True)


def page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])


def setup_document():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(1.0)
    sec.right_margin = Inches(1.0)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for level, size, color, before, after in [
        (1, 16, BLUE, 18, 10),
        (2, 13, BLUE, 14, 7),
        (3, 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = sec.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hr = header.add_run("HPDI低温泵声音工程对比 | 东德基准 vs 富瑞初始")
    set_run_font(hr, size=8.5, color="64748B")
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fr = footer.add_run("第 ")
    set_run_font(fr, size=8.5, color="64748B")
    page_field(footer)
    fr2 = footer.add_run(" 页")
    set_run_font(fr2, size=8.5, color="64748B")
    return doc


def main():
    with (OUT / "baseline_comparison.csv").open(encoding="utf-8-sig") as f:
        rows = list(csv.DictReader(f))
    by_speed = {int(r["speed_rpm"]): r for r in rows}

    doc = setup_document()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("声音工程对比基准报告")
    set_run_font(r, size=24, bold=True, color=INK)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(14)
    r = p2.add_run("HPDI低温泵 | 东德竞品基准与富瑞初始状态")
    set_run_font(r, size=13, bold=True, color=TEAL)

    meta = doc.add_table(rows=3, cols=2)
    set_table_geometry(meta, [1800, 7560])
    for i, (label, value) in enumerate([
        ("比较对象", "东德泵 2026-07-07 / 富瑞泵初始状态 2026-07-09"),
        ("比较原则", "相同转速、相同手机、相同录制角度和距离"),
        ("报告版本", "基准版 V2 | 2026-07-23"),
    ]):
        shade(meta.cell(i, 0), LIGHT_GRAY)
        set_cell_text(meta.cell(i, 0), label, bold=True, color=DARK_BLUE)
        set_cell_text(meta.cell(i, 1), value)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    add_callout(doc, "核心结论", "富瑞初始状态尚未在全部转速稳定优于东德。800 rpm 表现较好；700 rpm 的常见声略低，但偏响周期仍略高；900 rpm 和 1125 rpm 明显偏响，是后续优化重点。", GOLD)

    add_heading(doc, "1. 客户一眼能看懂的结果", 1)
    table = doc.add_table(rows=1, cols=6)
    headers = ["转速", "常见咚声", "较响咚声", "超过东德较响门槛", "判断", "建议"]
    widths = [900, 1500, 1500, 1500, 1560, 2400]
    for i, text in enumerate(headers):
        shade(table.cell(0, i), DARK_BLUE)
        set_cell_text(table.cell(0, i), text, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER, size=9)
    summaries = {
        700: ("低于东德 1.2 dB", "高于东德 1.8 dB", "11.5%", "接近基准", "降低偶发较响周期"),
        800: ("低于东德 7.9 dB", "低于东德 2.3 dB", "5.3%", "优于东德", "保持并做重复验证"),
        900: ("高于东德 7.6 dB", "高于东德 5.1 dB", "39.6%", "明显偏响", "最高优先级优化"),
        1125: ("高于东德 4.1 dB", "高于东德 3.4 dB", "47.2%", "明显偏响", "优化后延长录音复测"),
    }
    for speed, values in summaries.items():
        cells = table.add_row().cells
        fill = GREEN if "优于" in values[3] else RED if "明显" in values[3] else GOLD
        for i, text in enumerate([f"{speed} rpm", *values]):
            if i == 4:
                shade(cells[i], fill)
            set_cell_text(cells[i], text, bold=(i in [0, 4]), align=WD_ALIGN_PARAGRAPH.CENTER if i < 5 else WD_ALIGN_PARAGRAPH.LEFT, size=8.7)
    set_table_geometry(table, widths)
    add_para(doc, "这里的“低于/高于东德”是声音差值的方向，不是负分贝。所有客户可见声级均使用正数的估算 dB(A)。", bold_label="这里的“低于/高于东德”")

    add_heading(doc, "2. 三个指标分别代表什么", 1)
    defs = doc.add_table(rows=3, cols=2)
    set_table_geometry(defs, [2100, 7260])
    definitions = [
        ("常见咚声", "代表大多数正常工作周期中经常听到的咚声水平，用来判断日常听感。"),
        ("较响咚声", "代表一段测试里偏响的那部分周期，用来判断偶发冲击是否刺耳。"),
        ("超过门槛的比例", "东德自身约有 10% 的咚声会进入“较响”范围；富瑞比例越高，说明偏响周期越频繁。"),
    ]
    for i, (label, text) in enumerate(definitions):
        shade(defs.cell(i, 0), LIGHT_BLUE)
        set_cell_text(defs.cell(i, 0), label, bold=True, color=DARK_BLUE)
        set_cell_text(defs.cell(i, 1), text)

    doc.add_page_break()
    add_heading(doc, "3. 整段咚声趋势", 1)
    add_para(doc, "每个点代表一次有效咚声，黑线是连续 5 次咚声的趋势。它能直接显示你提到的“连续几个周期都较小，后面几个周期又变大”的现象。排气尖声和台架稳态背景不进入咚声统计。")
    add_figure(doc, PLOTS / "02_同转速包络直观对比.png", "图1  东德基准与富瑞初始状态的整段咚声趋势")
    add_callout(doc, "判读", "700 rpm 两者总体接近；800 rpm 富瑞多数周期较低；900 rpm 富瑞整段趋势整体上移，偏响不是单个偶发点，而是持续性差异。", LIGHT_BLUE)

    doc.add_page_break()
    add_heading(doc, "4. 常见咚声和较响咚声单独比较", 1)
    add_para(doc, "下图从每段录音中分别抽取一声具有代表性的常见咚声和较响咚声。曲线保留实际声级差异，没有人为把两条曲线峰值拉到同样高度，因此高低差可以直接比较；横轴 0 ms 附近为咚声峰值。")
    add_figure(doc, PLOTS / "03_大小咚代表事件叠加.png", "图2  常见咚声与较响咚声的单次曲线叠加")
    add_callout(doc, "最明显差异", "900 rpm 下，富瑞常见咚声和较响咚声均明显高于东德；800 rpm 下，富瑞较响咚声峰值低于东德。", LIGHT_BLUE)

    doc.add_page_break()
    add_heading(doc, "5. 基准评分与后续优化方向", 1)
    add_figure(doc, PLOTS / "04_一页式基准评分卡.png", "图3  富瑞相对东德的差值大小和偏响周期占比")
    add_para(doc, "第一优先级：900 rpm。常见咚声高 7.6 dB、较响咚声高 5.1 dB，且 39.6% 的事件超过东德较响门槛。", bold_label="第一优先级：")
    add_para(doc, "第二优先级：1125 rpm。两类咚声均偏高，超过门槛的比例达到 47.2%；现有竞品录音较短，优化后应延长录音复核。", bold_label="第二优先级：")
    add_para(doc, "保持项：800 rpm。当前结果优于东德，但需要至少再做两次相同工况重复测试，确认不是录音自动增益或单次工况波动造成。", bold_label="保持项：")
    add_para(doc, "改善项：700 rpm。常见声已经略低于东德，重点不是整体再降很多，而是压低少数偏响周期。", bold_label="改善项：")

    add_heading(doc, "6. 咚声可能出现在哪个动作阶段", 1)
    add_para(doc, "现有 1000 rpm 声音与功能数据的候选对齐显示，咚声主要集中在低温泵出口压力上升/打压开始，以及压力下降/泄压回程两个边界。统计中约 34 次位于压力上升段、30 次位于下降段，明显多于高压平台和低压等待阶段。")
    add_callout(doc, "工程判断", "大小咚声目前更像同一套换向和打压动作在不同周期中的冲击强弱变化，而不是两个始终固定、规律出现的独立声源。可能关联换向冲击、压力快速建立、行程端部阻力或卸压过程。", GOLD)
    add_para(doc, "该定位仍属于候选判断，因为视频和台架数据没有硬件同步触发。后续若要精确定位到换向阀、油缸端部或低温端阀动作，需要同步记录油缸位移、换向信号、系统油压和低温泵出口压力。")

    add_heading(doc, "7. 后续测试如何持续加入", 1)
    workflow = doc.add_table(rows=5, cols=3)
    set_table_geometry(workflow, [800, 2500, 6060])
    for i, text in enumerate(["步骤", "动作", "固定要求"]):
        shade(workflow.cell(0, i), DARK_BLUE)
        set_cell_text(workflow.cell(0, i), text, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    steps = [
        ("1", "录制", "同一手机和录制模式；手机、泵、分贝仪位置保持一致；每个稳定工况建议不少于 60 s。"),
        ("2", "登记", "记录测试日期、油缸/密封/程序等配置、转速、压力和异常现象，形成唯一测试编号。"),
        ("3", "分析", "按同一频带、同一背景扣除和同一排气污染规则提取整体包络、常见咚声和较响咚声。"),
        ("4", "比较", "正式结论只与同转速东德基准比较；不同转速只能观察趋势，不能直接判断产品优劣。"),
    ]
    for step in steps:
        cells = workflow.add_row().cells
        for i, text in enumerate(step):
            set_cell_text(cells[i], text, bold=(i == 0), align=WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT, size=9)
    add_para(doc, "配套 Excel 已预留 6 个对比槽位。选择测试编号后，整体包络、常见咚声、较响咚声和结果表会同步切换；未选测试不进入曲线。")

    add_heading(doc, "8. 计算口径与使用边界", 1)
    add_para(doc, "声级单位：视频中的 Smart Sensor AR824 屏幕可识别为 dBA。由于 FAST/SLOW 时间计权标志看不清，手机又可能启用自动增益，因此报告中的绝对 dB(A)为工程估算值；它主要帮助阅读曲线，不作为认证或法规声级。", bold_label="声级单位：")
    add_para(doc, "可靠主指标：同一录制条件下“富瑞比东德高/低多少 dB”。固定校准偏移会在相减时抵消，因此相对差值比绝对声级更可靠。", bold_label="可靠主指标：")
    add_para(doc, "干扰处理：在 45-650 Hz 范围内提取往复冲击；每个事件扣除前后局部台架背景；高频排气尖声作为污染标记并剔除。算法只能降低干扰，不能证明环境声被百分之百分离。", bold_label="干扰处理：")
    add_para(doc, "转速原则：当前数据已显示明显转速依赖，因此东德基准按 700/800/900/1125 rpm 分开保存。后续只有在重复试验证明各转速差异长期小于约 1 dB 后，才考虑合并。", bold_label="转速原则：")

    doc.save(DOCX)
    print(DOCX)


if __name__ == "__main__":
    main()
